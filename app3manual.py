import streamlit as st
import pandas as pd
from docx import Document
import gspread
from google.oauth2.service_account import Credentials
import io
import requests
import openpyxl
import time

# --- 1. CONFIG ---
st.set_page_config(page_title="Engineer Master Hybrid v7.17", layout="wide")

TOKEN = st.secrets["TOKEN"]
ID_USER_WORD = st.secrets.get("CHAT_ID_1") 
ID_USER_FULL = st.secrets.get("CHAT_ID_2")
GSHEET_URL = st.secrets["GSHEET_URL"]

# --- DAFTAR 36 AITEM UNTUK DETEKSI TEKS ---
LIST_AITEM = [
    "Seiring waktu, saya bisa memaklumi kesalahan pribadi yang pernah dilakukan.",
    "Ketika membuat kesalahan, saya fokus pada perbaikan daripada terus menerus menyalahkan diri sendiri.",
    "Saya memilih untuk berdamai dengan kekurangan diri sendiri.",
    "Sulit bagi saya untuk berhenti menyalahkan diri sendiri.",
    "Muncul perasaan benci ketika saya mengingat kesalahan diri sendiri.",
    "Saya terjebak dalam penyesalan atas kegagalan diri sendiri.",
    "Pikiran negatif tentang diri sendiri mulai memudar seiring waktu.",
    "Saya dapat memahami diri sendiri atas kesalahan yang telah saya lakukan.",
    "Saat ingatan yang mengganggu tentang diri sendiri muncul, saya mampu melepaskannya.",
    "Sulit bagi saya untuk berhenti memikirkan hal-hal buruk yang pernah menimpa diri sendiri.",
    "Pikiran tentang kesalahan diri sendiri terus muncul walaupun sudah berusaha melupakannya.",
    "Saya sering susah berkonsentrasi karena teringat pada kesalahan diri sendiri yang telah lalu.",
    "Saya dapat memaklumi bahwa setiap orang pasti pernah melakukan kekeliruan.",
    "Saya mencoba memahami alasan dibalik tindakan orang lain yang telah menyakiti saya.",
    "Saya menyadari bahwa ada alasan tertentu yang membuat orang lain sulit untuk bertindak benar.",
    "Memandang orang yang menyakiti saya sebagai pribadi yang memiliki karakter buruk.",
    "Saya tidak bisa menerima alasan apapun dari orang yang telah mengecewakan saya.",
    "Sangat sulit bagi saya untuk mengerti mengapa seseorang berbuat jahat kepada saya.",
    "Pikiran buruk terhadap orang yang pernah menyakiti saya perlahan mulai menghilang.",
    "Saya merasa sudah tidak lagi menyimpan kebencian terhadap orang yang pernah menyakiti saya.",
    "Mudah bagi saya melepaskan rasa benci yang tertuju pada orang yang pernah berbuat salah.",
    "Saya terus membayangkan hal-hal negatif terjadi pada orang yang telah menyakiti saya.",
    "Sulit bagi saya untuk menghilangkan pandangan negatif terhadap orang yang pernah berbuat salah.",
    "Rasa kesal muncul kembali setiap kali saya mengingat perlakuan orang yang menyakiti saya.",
    "Seiring berjalannya waktu, saya mulai bisa menerima kenyataan pahit yang terjadi dalam hidup dengan lapang dada.",
    "Saya sadar untuk tidak menyalahkan nasib atas kejadian buruk yang menimpa.",
    "Mampu menerima kenyataan bahwa hidup tidak selalu berjalan sesuai dengan rencana saya.",
    "Saya merasa semesta tidak adil karena terus memberikan cobaan yang berat.",
    "Sering merasa terjebak dalam nasib buruk yang seolah-olah tidak pernah berakhir di hidup saya.",
    "Terus-menerus mengeluhkan nasib buruk yang menimpa diri saya menjadi hal yang sulit untuk dihentikan.",
    "Pikiran tentang kejadian buruk di masa lalu tidak lagi mengganggu saya untuk berkonsentrasi sehari-hari.",
    "Saya merasa sudah bisa berdamai dengan bayangan tentang masa-masa sulit yang pernah dialami.",
    "Saya mampu mengalihkan fokus dari peristiwa yang mengecewakan ke hal-hal yang lebih produktif.",
    "Sangat sulit bagi saya untuk tidak memikirkan kegagalan yang pernah dialami.",
    "Saya merasa terjebak dalam memori tentang kejadian buruk yang pernah saya alami.",
    "Bayangan mengenai ketidakadilan hidup di masa lalu sering kali muncul tanpa bisa saya kendalikan."
]

# --- 2. CORE FUNCTIONS ---

def get_gsheet_client():
    scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scope)
    return gspread.authorize(creds)

def find_first_empty_row(ws):
    col_c = ws.col_values(3) 
    for r in range(4, 34): 
        if r > len(col_c) or not col_c[r-1]:
            return r
    return 34

def proses_excel_cvi():
    try:
        client = get_gsheet_client(); ss = client.open_by_url(GSHEET_URL)
        wb = openpyxl.load_workbook("CVI Aiken Zuyy.xlsx")
        for s_name in ["KEJELASAN", "RELEVANSI", "KESESUAIAN"]:
            ws_gs = ss.worksheet(s_name)
            all_vals = ws_gs.get_all_values()
            ws_xl = wb[s_name]
            for idx, row_data in enumerate(all_vals[3:]):
                target_row = 4 + idx
                if target_row > 33 or len(row_data) < 2 or not row_data[1]: break
                ws_xl.cell(row=target_row, column=2, value=row_data[1])
                for col_idx, val in enumerate(row_data[2:38]):
                    try: ws_xl.cell(row=target_row, column=3+col_idx, value=int(float(val)))
                    except: ws_xl.cell(row=target_row, column=3+col_idx, value=0)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0); return buf
    except: return None

def send_tele(chat_id, file_buf, fname, caption):
    url = f"https://api.telegram.org/bot{TOKEN}/sendDocument"
    file_buf.seek(0)
    return requests.post(url, data={'chat_id': chat_id, 'caption': caption}, files={'document': (fname, file_buf)})

# --- 3. UI ENGINE ---

st.title("🖥️ MASTER HYBRID v7.17 (TEXT DETECTION MODE)")

if 'batch_data' not in st.session_state:
    st.session_state.batch_data = None

up_file = st.file_uploader("Upload Excel Aiken", type=["xlsx"])

if up_file:
    if st.button("⚙️ PROCESS DATA"):
        try:
            xl = pd.ExcelFile(up_file); temp_db = {}
            for s_name in ["KEJELASAN", "RELEVANSI", "KESESUAIAN"]:
                raw_xl = pd.read_excel(up_file, sheet_name=s_name, header=None, skiprows=3)
                raw_xl = raw_xl.reindex(columns=range(40), fill_value=0)
                rows = []
                for i in range(len(raw_xl)):
                    name = raw_xl.iloc[i, 1] 
                    if pd.isna(name) or str(name).strip() == "": continue
                    job = raw_xl.iloc[i, 2] if s_name == "KEJELASAN" else ""
                    start_skor = 3 if s_name == "KEJELASAN" else 2
                    row_dict = {"Nama": str(name), "Pekerjaan": str(job) if not pd.isna(job) else ""}
                    for a_idx in range(36):
                        val = raw_xl.iloc[i, start_skor + a_idx]
                        try: row_dict[f"A{a_idx+1}"] = int(val) if not pd.isna(val) else 0
                        except: row_dict[f"A{a_idx+1}"] = 0
                    rows.append(row_dict)
                temp_db[s_name] = pd.DataFrame(rows)
            st.session_state.batch_data = temp_db
            st.success("Excel Terbaca! Siap Deteksi Teks."); st.rerun()
        except Exception as e: st.error(f"Error: {e}")

if st.session_state.batch_data:
    tabs = st.tabs(["📊 KJ", "📈 REL", "📉 KES"])
    sheets = ["KEJELASAN", "RELEVANSI", "KESESUAIAN"]
    for i, tab in enumerate(tabs):
        with tab:
            config = {"Pekerjaan": None} if sheets[i] != "KEJELASAN" else {}
            st.session_state.batch_data[sheets[i]] = st.data_editor(st.session_state.batch_data[sheets[i]], num_rows="dynamic", key=f"ed_{sheets[i]}", column_config=config)

    if st.button("🚀 EXECUTE SMART TEXT SYNC"):
        try:
            with st.spinner("Detecting Text in Word Table..."):
                client = get_gsheet_client(); ss = client.open_by_url(GSHEET_URL)
                w_tmpl = "Form Validasi Expert Judgement Ayinn Ver. 3.docx"
                
                db_kj = st.session_state.batch_data["KEJELASAN"]
                db_rel = st.session_state.batch_data["RELEVANSI"]
                db_kes = st.session_state.batch_data["KESESUAIAN"]

                for idx in range(len(db_kj)):
                    name = str(db_kj.iloc[idx]["Nama"])
                    job = str(db_kj.iloc[idx]["Pekerjaan"]) 
                    
                    # 1. Update GSheets
                    for s_name in sheets:
                        ws = ss.worksheet(s_name)
                        target_row = find_first_empty_row(ws)
                        if target_row <= 33:
                            curr_df = st.session_state.batch_data[s_name]
                            skor = [int(curr_df.iloc[idx][f"A{k+1}"]) for k in range(36)]
                            ws.update_cell(target_row, 2, name)
                            cells = ws.range(target_row, 3, target_row, 3 + len(skor) - 1)
                            for s_idx, v in enumerate(skor): cells[s_idx].value = v
                            ws.update_cells(cells)

                    # 2. Injeksi Word (LOGIKA DETEKSI TEKS PERSIS PORTAL)
                    doc = Document(w_tmpl)
                    for p in doc.paragraphs:
                        if "Nama\t\t:" in p.text: p.text = f"Nama\t\t: {name}"
                        if "Pekerjaan\t:" in p.text: p.text = f"Pekerjaan\t: {job}"
                    
                    if doc.tables:
                        table = doc.tables[0]
                        skj = [int(db_kj.iloc[idx][f"A{k+1}"]) for k in range(36)]
                        srel = [int(db_rel.iloc[idx][f"A{k+1}"]) for k in range(36)]
                        skes = [int(db_kes.iloc[idx][f"A{k+1}"]) for k in range(36)]
                        
                        for row in table.rows:
                            # Ambil teks aitem di Word, hilangkan spasi & kecilkan huruf
                            word_text = "".join(row.cells[2].text.split()).lower()
                            
                            # Bandingkan dengan 36 List Aitem Master
                            for a_idx, master_text in enumerate(LIST_AITEM):
                                master_clean = "".join(master_text.split()).lower()[:60]
                                if master_clean in word_text:
                                    # Match Ketemu! Suntik Skor
                                    row.cells[3].text = str(skj[a_idx])
                                    row.cells[4].text = str(srel[a_idx])
                                    row.cells[5].text = str(skes[a_idx])
                                    break # Keluar loop master, lanjut ke baris tabel berikutnya
                    
                    word_buf = io.BytesIO(); doc.save(word_buf); word_buf.seek(0)
                    send_tele(ID_USER_WORD, word_buf, f"Form Validasi Expert Judgement Forgiveness_{name}.docx", f"✅ Word: {name}")
                    word_buf.seek(0); send_tele(ID_USER_FULL, word_buf, f"Form Validasi Expert Judgement Forgiveness_{name}.docx", f"✅ Full Log: {name}")

                # 3. KIRIM EXCEL REKAP KUMULATIF KE ADMIN
                excel_rekap = proses_excel_cvi()
                if excel_rekap:
                    send_tele(ID_USER_FULL, excel_rekap, f"Rekap_Cumulative_{int(time.time())}.xlsx", "📊 REKAP TOTAL DATABASE")
                
                st.success("Selesai! Deteksi Teks Berhasil & Aitem 36 Aman."); st.session_state.batch_data = None
        except Exception as e: st.error(f"Pipeline Error: {e}")
